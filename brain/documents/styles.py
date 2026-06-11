"""IEEE/ISO-compliant document styles for python-docx."""
from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn


def apply_base_styles(doc: Document) -> None:
    """Apply professional IEEE/ISO-compliant styles to the document."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        hf = heading_style.font
        hf.name = "Calibri"
        hf.bold = True
        hf.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        if level == 1:
            hf.size = Pt(18)
            heading_style.paragraph_format.space_before = Pt(24)
            heading_style.paragraph_format.space_after = Pt(12)
        elif level == 2:
            hf.size = Pt(14)
            heading_style.paragraph_format.space_before = Pt(18)
            heading_style.paragraph_format.space_after = Pt(8)
        else:
            hf.size = Pt(12)
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_header_footer(doc: Document, title: str, project_name: str, version: str) -> None:
    """Add header with project info and footer with page numbers."""
    section = doc.sections[0]

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = f"{project_name}  |  {title}  |  v{version}"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.runs[0] if hp.runs else hp.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Footer with page number
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fld_char_begin = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    run._r.append(fld_char_begin)
    run2 = fp.add_run()
    instr = run2._r.makeelement(qn("w:instrText"), {})
    instr.text = " PAGE "
    run2._r.append(instr)
    run3 = fp.add_run()
    fld_char_end = run3._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run3._r.append(fld_char_end)
    for r in fp.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def add_cover_page(doc: Document, title: str, project_name: str, version: str, standard: str) -> None:
    """Add a professional cover page."""
    for _ in range(4):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(project_name.upper())
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Versione {version}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    if standard and standard != "minimal":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label = {"ieee830": "IEEE 830 / ISO 29148", "iso29148": "ISO/IEC/IEEE 29148:2018"}.get(standard, standard)
        run = p.add_run(f"Standard: {label}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)


def add_code_block(doc: Document, code: str, language: str = "") -> None:
    """Add a monospace code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    if language:
        label_run = p.add_run(f"[{language}]\n")
        label_run.font.size = Pt(8)
        label_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
