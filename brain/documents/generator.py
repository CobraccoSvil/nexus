"""Document generator: creates .docx files from structured JSON content."""
from __future__ import annotations

import json
import logging
import os
from datetime import datetime
from pathlib import Path

from docx import Document

from brain.documents.styles import (
    add_code_block,
    add_cover_page,
    add_header_footer,
    add_table,
    apply_base_styles,
)
from brain.documents.templates import get_template

logger = logging.getLogger(__name__)


def _as_text(value) -> str:
    """Coerce un valore `content` arbitrario a stringa renderizzabile.

    FIX 5 (anti-malformazione): il modello docs_generator a volte annida liste
    o oggetti dentro `content` invece di una stringa. Prima `content.strip()` e
    `len(content)` sollevavano eccezioni non gestite (AttributeError/TypeError)
    durante il rendering, facendo fallire l'intera generazione o producendo un
    .docx corrotto. Qui normalizziamo sempre a testo leggibile.
    """
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, (list, tuple)):
        return "\n".join(_as_text(v) for v in value if v is not None)
    if isinstance(value, dict):
        for key in ("content", "text", "value"):
            if key in value:
                return _as_text(value[key])
        return json.dumps(value, ensure_ascii=False, indent=2)
    return str(value)


class DocumentGenerator:
    """Generates professional .docx documents from structured JSON content."""

    def generate(
        self,
        doc_type: str,
        content_json: str,
        output_path: str,
        standard: str = "ieee830",
        title: str = "",
        project_name: str = "",
    ) -> dict:
        """Generate a .docx document.

        Args:
            doc_type: One of functional_analysis, technical_analysis, er_diagram,
                      project_management, release_notes
            content_json: JSON string with {"sections": [...]} structure
            output_path: Absolute path for the output .docx file
            standard: Document standard (ieee830, iso29148, minimal)
            title: Document title (falls back to template default)
            project_name: Project name for header/cover

        Returns:
            dict with file_path, page_count, section_count, error
        """
        try:
            content = json.loads(content_json) if isinstance(content_json, str) else content_json
        except json.JSONDecodeError as e:
            return {"file_path": "", "page_count": 0, "section_count": 0, "error": f"JSON non valido: {e}"}

        template = get_template(doc_type)
        title = title or template.get("title_default", "Documento")
        standard = standard or template.get("standard", "minimal")
        version = content.get("version", "1.0.0")
        sections = content.get("sections", [])

        # If no sections provided, use template skeleton
        if not sections:
            sections = template.get("sections", [])

        doc = Document()
        apply_base_styles(doc)
        add_cover_page(doc, title, project_name, version, standard)
        add_header_footer(doc, title, project_name, version)

        section_count = 0
        for section in sections:
            section_count += self._render_section(doc, section, level=1)

        # Ensure output directory exists
        out_path = Path(output_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)

        # Write to temp file first, then rename (atomic)
        tmp_path = out_path.with_suffix(".tmp.docx")
        try:
            doc.save(str(tmp_path))
            if out_path.exists():
                out_path.unlink()
            tmp_path.rename(out_path)
        except Exception as e:
            tmp_path.unlink(missing_ok=True)
            return {"file_path": "", "page_count": 0, "section_count": 0, "error": f"Errore salvataggio: {e}"}

        # Estimate page count (~3000 chars per page)
        total_chars = sum(len(_as_text(s.get("content", ""))) for s in sections)
        page_count = max(1, total_chars // 3000 + 1)

        logger.info("Documento generato: %s (%d sezioni, ~%d pagine)", output_path, section_count, page_count)

        return {
            "file_path": str(out_path),
            "page_count": page_count,
            "section_count": section_count,
            "error": "",
        }

    def _render_section(self, doc: Document, section: dict, level: int) -> int:
        """Render a section and its subsections recursively. Returns count."""
        number = section.get("number", "")
        title = section.get("title", "")
        content = _as_text(section.get("content", ""))
        subsections = section.get("subsections", [])
        if not isinstance(subsections, list):
            subsections = []

        heading_level = min(level, 3)
        heading_text = f"{number}. {title}" if number else title
        doc.add_heading(heading_text, level=heading_level)

        count = 1

        if content:
            self._render_content(doc, content)

        # Render tables if present
        if "table" in section:
            tbl = section["table"]
            headers = tbl.get("headers", [])
            rows = tbl.get("rows", [])
            if headers and rows:
                add_table(doc, headers, rows)

        # Render code blocks if present
        if "code" in section:
            code_data = section["code"]
            if isinstance(code_data, dict):
                add_code_block(doc, code_data.get("content", ""), code_data.get("language", ""))
            elif isinstance(code_data, str):
                add_code_block(doc, code_data)

        # Render subsections
        for sub in subsections:
            count += self._render_section(doc, sub, level + 1)

        return count

    def _render_content(self, doc: Document, content: str) -> None:
        """Render content text, handling basic formatting."""
        lines = content.strip().split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # Bullet points
            if stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(stripped[2:], style="List Bullet")
            # Numbered list
            elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in (".", ")"):
                p = doc.add_paragraph(stripped[2:].strip(), style="List Number")
            # Mermaid code block
            elif stripped.startswith("```"):
                # Collect until closing ```
                continue
            else:
                doc.add_paragraph(stripped)
